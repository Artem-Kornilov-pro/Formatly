import io
from pathlib import Path

from docx import Document

from app.pipeline.parser import parse_docx


def _build_docx(tmp_path: Path) -> Path:
    document = Document()
    document.add_heading("Introduction", level=1)
    document.add_paragraph("")  # blank spacer, should be skipped
    document.add_paragraph("This is the first body paragraph.")
    document.add_heading("Background", level=2)
    document.add_paragraph("This is the second body paragraph.")

    path = tmp_path / "input.docx"
    document.save(str(path))
    return path


def test_parse_docx_extracts_text_and_style_skipping_blank_paragraphs(tmp_path: Path):
    path = _build_docx(tmp_path)

    paragraphs = parse_docx(path)

    texts = [p.text for p in paragraphs]
    assert texts == [
        "Introduction",
        "This is the first body paragraph.",
        "Background",
        "This is the second body paragraph.",
    ]
    assert paragraphs[0].style_name == "Heading 1"
    assert paragraphs[3].style_name == "Normal"


def test_parse_docx_preserves_original_paragraph_index(tmp_path: Path):
    path = _build_docx(tmp_path)

    paragraphs = parse_docx(path)

    # index 1 is the blank spacer paragraph that got dropped, so the
    # remaining indices must keep the gap rather than being renumbered -
    # the formatter later maps back onto document.paragraphs by index.
    assert [p.index for p in paragraphs] == [0, 2, 3, 4]


def test_parse_docx_accepts_in_memory_generated_file(tmp_path: Path):
    document = Document()
    document.add_paragraph("Hello, Formatly!")
    buffer = io.BytesIO()
    document.save(buffer)

    path = tmp_path / "roundtrip.docx"
    path.write_bytes(buffer.getvalue())

    paragraphs = parse_docx(path)

    assert len(paragraphs) == 1
    assert paragraphs[0].text == "Hello, Formatly!"
