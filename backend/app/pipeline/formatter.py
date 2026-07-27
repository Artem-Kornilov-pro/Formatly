from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from app.pipeline.rules import FormattingRules
from app.pipeline.schemas import (
    TOC_HEADING_TEXT,
    ClassifiedParagraph,
    ParagraphRole,
    should_insert_generated_title,
)

_HEADING_ROLES = {ParagraphRole.HEADING_1, ParagraphRole.HEADING_2, ParagraphRole.HEADING_3}
# heading_size_bump_pt scaled by level: heading 1 gets the full bump, heading
# 2 gets half, heading 3 gets none.
_HEADING_BUMP_SCALE = {
    ParagraphRole.HEADING_1: 1.0,
    ParagraphRole.HEADING_2: 0.5,
    ParagraphRole.HEADING_3: 0.0,
}
_ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
_OUTLINE_LEVELS = {
    ParagraphRole.HEADING_1: 0,
    ParagraphRole.HEADING_2: 1,
    ParagraphRole.HEADING_3: 2,
}

_TOC_FIELD_CODE = ' TOC \\o "1-3" \\h \\z \\u '
_TOC_PLACEHOLDER_TEXT = (
    "Автособираемое оглавление появится здесь при открытии документа в Word "
    "(или после обновления поля вручную: клавиша F9)."
)
# w:updateFields is a ZeroOrOne element in the w:settings schema, but
# python-docx doesn't expose a getter/setter for it, so it has to be spliced
# in by hand via insert_element_before. This is the ordered list of sibling
# tags the OOXML schema requires to come after it - python-docx keeps the
# equivalent list private (CT_Settings._tag_seq, deleted off the class right
# after it's defined), so it's copied here rather than depended on.
_SETTINGS_TAGS_AFTER_UPDATE_FIELDS = (
    "w:hdrShapeDefaults",
    "w:footnotePr",
    "w:endnotePr",
    "w:compat",
    "w:docVars",
    "w:rsids",
    "m:mathPr",
    "w:attachedSchema",
    "w:themeFontLang",
    "w:clrSchemeMapping",
    "w:doNotIncludeSubdocsInStats",
    "w:doNotAutoCompressPictures",
    "w:forceUpgrade",
    "w:captions",
    "w:readModeInkLockDown",
    "w:smartTagType",
    "sl:schemaLibrary",
    "w:shapeDefaults",
    "w:doNotEmbedSmartTags",
    "w:decimalSymbol",
    "w:listSeparator",
)


def _set_run_font(run, font_family: str) -> None:
    run.font.name = font_family
    # python-docx's font.name setter only writes w:ascii - Cyrillic and other
    # non-Latin runs are rendered via w:hAnsi/w:eastAsia/w:cs, so without this
    # the font change silently doesn't apply to Russian text.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:eastAsia", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), font_family)


def apply_formatting(
    input_path: Path,
    output_path: Path,
    classified: list[ClassifiedParagraph],
    rules: FormattingRules,
    generated_title: str | None = None,
) -> list[str]:
    document = Document(str(input_path))

    for section in document.sections:
        section.top_margin = Mm(rules.margins_mm.top)
        section.bottom_margin = Mm(rules.margins_mm.bottom)
        section.left_margin = Mm(rules.margins_mm.left)
        section.right_margin = Mm(rules.margins_mm.right)

    role_by_index = {paragraph.index: paragraph.role for paragraph in classified}
    completion_by_index = (
        {paragraph.index: paragraph.completion for paragraph in classified if paragraph.completion}
        if rules.ai_light_editing_enabled
        else {}
    )
    body_alignment = _ALIGNMENTS[rules.paragraph_alignment]

    for index, paragraph in enumerate(document.paragraphs):
        role = role_by_index.get(index, ParagraphRole.BODY)
        is_heading = role in _HEADING_ROLES

        size_pt = rules.font_size_pt
        if is_heading:
            size_pt += round(rules.heading_size_bump_pt * _HEADING_BUMP_SCALE[role])

        paragraph.paragraph_format.line_spacing = rules.line_spacing

        if is_heading:
            paragraph.paragraph_format.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if rules.center_headings else WD_ALIGN_PARAGRAPH.LEFT
            )
            paragraph.paragraph_format.first_line_indent = Mm(0)
            if role == ParagraphRole.HEADING_1:
                paragraph.paragraph_format.page_break_before = rules.page_break_before_heading_1
            if rules.generate_toc:
                # Our headings are assigned by content classification, not
                # necessarily Word's built-in Heading 1/2/3 styles, so the
                # TOC field below needs an explicit outline level to find
                # them when it's (re)generated.
                _set_outline_level(paragraph, _OUTLINE_LEVELS[role])
        else:
            paragraph.paragraph_format.alignment = body_alignment
            paragraph.paragraph_format.first_line_indent = (
                Mm(rules.paragraph_indent_mm) if rules.paragraph_indent_enabled else Mm(0)
            )

        completion = completion_by_index.get(index)
        if completion:
            _append_completion(paragraph, completion)

        for run in paragraph.runs:
            _set_run_font(run, rules.font_family)
            run.font.size = Pt(size_pt)
            run.font.bold = is_heading and rules.bold_headings
            run.font.italic = is_heading and rules.italic_headings

    inserted_title = (
        generated_title
        if rules.ai_light_editing_enabled and should_insert_generated_title(classified, generated_title)
        else None
    )
    if inserted_title:
        _insert_generated_title_paragraph(document, rules, inserted_title)
        # every paragraph that existed before the title was spliced in front
        # of it just shifted down by one position
        role_by_index = {index + 1: role for index, role in role_by_index.items()}
        role_by_index[0] = ParagraphRole.TITLE

    toc_inserted = rules.generate_toc and _insert_table_of_contents(document, role_by_index, rules)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))

    return _describe_changes(rules, toc_inserted, inserted_title, len(completion_by_index))


def _append_completion(paragraph, completion: str) -> None:
    separator = "" if not paragraph.text or paragraph.text[-1].isspace() else " "
    paragraph.add_run(separator + completion)


def _insert_generated_title_paragraph(
    document: Document, rules: FormattingRules, generated_title: str
) -> None:
    existing_paragraphs = document.paragraphs
    paragraph = document.add_paragraph()
    run = paragraph.add_run(generated_title)
    _set_run_font(run, rules.font_family)
    run.font.size = Pt(rules.font_size_pt + rules.heading_size_bump_pt)
    run.font.bold = rules.bold_headings
    run.font.italic = rules.italic_headings
    paragraph.paragraph_format.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if rules.center_headings else WD_ALIGN_PARAGRAPH.LEFT
    )
    paragraph.paragraph_format.first_line_indent = Mm(0)
    if existing_paragraphs:
        existing_paragraphs[0]._p.addprevious(paragraph._p)


def _set_outline_level(paragraph, level: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    outline_lvl = ppr.find(qn("w:outlineLvl"))
    if outline_lvl is None:
        outline_lvl = OxmlElement("w:outlineLvl")
        ppr.append(outline_lvl)
    outline_lvl.set(qn("w:val"), str(level))


def _add_field_char(paragraph, fld_char_type: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), fld_char_type)
    run._r.append(fld_char)


def _add_instr_text(paragraph, instr_text: str) -> None:
    run = paragraph.add_run()
    instr_elm = OxmlElement("w:instrText")
    instr_elm.set(qn("xml:space"), "preserve")
    instr_elm.text = instr_text
    run._r.append(instr_elm)


def _build_toc_heading_paragraph(document: Document, rules: FormattingRules):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(TOC_HEADING_TEXT)
    _set_run_font(run, rules.font_family)
    run.font.size = Pt(rules.font_size_pt + rules.heading_size_bump_pt)
    run.font.bold = rules.bold_headings
    run.font.italic = rules.italic_headings
    paragraph.paragraph_format.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if rules.center_headings else WD_ALIGN_PARAGRAPH.LEFT
    )
    paragraph.paragraph_format.first_line_indent = Mm(0)
    return paragraph._p


def _build_toc_field_paragraph(document: Document, rules: FormattingRules):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Mm(0)
    _add_field_char(paragraph, "begin")
    _add_instr_text(paragraph, _TOC_FIELD_CODE)
    _add_field_char(paragraph, "separate")
    placeholder = paragraph.add_run(_TOC_PLACEHOLDER_TEXT)
    _set_run_font(placeholder, rules.font_family)
    placeholder.font.size = Pt(rules.font_size_pt)
    placeholder.font.italic = True
    _add_field_char(paragraph, "end")
    return paragraph._p


def _enable_update_fields_on_open(document: Document) -> None:
    settings_elm = document.settings.element
    if settings_elm.find(qn("w:updateFields")) is not None:
        return
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings_elm.insert_element_before(update_fields, *_SETTINGS_TAGS_AFTER_UPDATE_FIELDS)


def _insert_table_of_contents(
    document: Document, role_by_index: dict[int, ParagraphRole], rules: FormattingRules
) -> bool:
    if not any(role in _HEADING_ROLES for role in role_by_index.values()):
        return False

    paragraphs = document.paragraphs
    title_indices = [i for i, role in role_by_index.items() if role == ParagraphRole.TITLE]
    anchor_index = max(title_indices) if title_indices else -1
    next_paragraph = paragraphs[anchor_index + 1] if anchor_index + 1 < len(paragraphs) else None

    heading_p = _build_toc_heading_paragraph(document, rules)
    field_p = _build_toc_field_paragraph(document, rules)

    if anchor_index >= 0:
        anchor = paragraphs[anchor_index]._p
        anchor.addnext(heading_p)
        heading_p.addnext(field_p)
    else:
        first = paragraphs[0]._p
        first.addprevious(heading_p)
        first.addprevious(field_p)

    # Start the main content on a fresh page after the TOC, matching how
    # a "СОДЕРЖАНИЕ" page conventionally stands on its own.
    if next_paragraph is not None:
        next_paragraph.paragraph_format.page_break_before = True

    _enable_update_fields_on_open(document)
    return True


def _describe_changes(
    rules: FormattingRules,
    toc_inserted: bool,
    inserted_title: str | None,
    completions_applied: int,
) -> list[str]:
    changes = [
        f"set page margins to {rules.margins_mm.top}/{rules.margins_mm.bottom}/"
        f"{rules.margins_mm.left}/{rules.margins_mm.right}mm (top/bottom/left/right)",
        f"applied {rules.font_family} {rules.font_size_pt}pt, line spacing {rules.line_spacing}, "
        f"{rules.paragraph_alignment} alignment to body text",
    ]

    if rules.paragraph_indent_enabled:
        changes.append(
            f"applied a {rules.paragraph_indent_mm}mm first-line indent to body paragraphs"
        )

    heading_style = []
    if rules.bold_headings:
        heading_style.append("bold")
    if rules.italic_headings:
        heading_style.append("italic")
    if rules.center_headings:
        heading_style.append("centered")
    if heading_style:
        changes.append(
            f"styled headings as {', '.join(heading_style)}, "
            f"enlarged by up to {rules.heading_size_bump_pt}pt"
        )

    if rules.page_break_before_heading_1:
        changes.append("inserted a page break before each top-level heading")

    if toc_inserted:
        changes.append(
            f'inserted an automatic table of contents ("{TOC_HEADING_TEXT}") that Word '
            "regenerates when the document is opened, starting the main content on a new page"
        )

    if inserted_title:
        changes.append(f'generated a document title ("{inserted_title}") since none was found')

    if completions_applied:
        changes.append(
            f"completed {completions_applied} paragraph(s) that looked cut off mid-sentence"
        )

    return changes
