import io
import uuid

from docx import Document
from sqlalchemy import select

from app.core.storage import input_file_path, output_file_path
from app.models.formatting_profile import FormattingProfile
from app.models.job import Job, JobStatus
from app.models.user import User
from app.models.validation_report import ValidationReport
from app.pipeline.run import process_job
from app.pipeline.schemas import ClassifiedParagraph, ParagraphRole
from tests.conftest import test_session_maker as make_test_session

RULES = {
    "font_family": "Times New Roman",
    "font_size_pt": 14,
    "line_spacing": 1.5,
    "margins_mm": {"top": 20, "bottom": 20, "left": 30, "right": 15},
}


class FakeClassifier:
    def __init__(self, roles: dict[int, ParagraphRole]):
        self._roles = roles

    def classify(self, paragraphs) -> list[ClassifiedParagraph]:
        return [
            ClassifiedParagraph(
                index=p.index, text=p.text, role=self._roles.get(p.index, ParagraphRole.BODY)
            )
            for p in paragraphs
        ]


def _docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Introduction", level=1)
    document.add_paragraph("Body paragraph.")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


async def _create_job(rules: dict = RULES) -> Job:
    async with make_test_session() as session:
        user = User(email=f"{uuid.uuid4()}@example.com", password_hash="not-a-real-hash")
        profile = FormattingProfile(name="Test profile", rules=rules)
        session.add_all([user, profile])
        await session.flush()

        job = Job(
            user_id=user.id,
            profile_id=profile.id,
            status=JobStatus.PENDING,
            input_file="thesis.docx",
        )
        session.add(job)
        await session.commit()
        await session.refresh(job)
        return job


async def test_process_job_runs_full_pipeline_and_marks_done():
    job = await _create_job()
    input_file_path(job.id).parent.mkdir(parents=True, exist_ok=True)
    input_file_path(job.id).write_bytes(_docx_bytes())

    fake_classifier = FakeClassifier({0: ParagraphRole.HEADING_1, 1: ParagraphRole.BODY})

    await process_job(job.id, classifier=fake_classifier, session_maker=make_test_session)

    async with make_test_session() as session:
        refreshed = await session.get(Job, job.id)
        assert refreshed.status == JobStatus.DONE
        assert refreshed.output_file == output_file_path(job.id).name
        assert refreshed.error_message is None

        report = await session.scalar(
            select(ValidationReport).where(ValidationReport.job_id == job.id)
        )
        assert report is not None
        assert report.issues_found == []
        assert len(report.issues_fixed) == 3

    assert output_file_path(job.id).exists()


async def test_process_job_marks_failed_and_records_error_on_exception():
    job = await _create_job()
    # deliberately no input file written, so parsing raises FileNotFoundError

    await process_job(job.id, classifier=FakeClassifier({}), session_maker=make_test_session)

    async with make_test_session() as session:
        refreshed = await session.get(Job, job.id)
        assert refreshed.status == JobStatus.FAILED
        assert refreshed.error_message
        assert refreshed.output_file is None


async def test_process_job_is_a_noop_for_unknown_job_id():
    # should not raise even though the job doesn't exist
    await process_job(uuid.uuid4(), classifier=FakeClassifier({}), session_maker=make_test_session)


async def test_reprocessing_after_a_failure_clears_the_old_error_message():
    job = await _create_job()
    # first attempt fails: no input file written yet
    await process_job(job.id, classifier=FakeClassifier({}), session_maker=make_test_session)

    async with make_test_session() as session:
        failed = await session.get(Job, job.id)
        assert failed.status == JobStatus.FAILED
        assert failed.error_message

    # fix the underlying problem and retry, same as re-queuing a failed job
    input_file_path(job.id).parent.mkdir(parents=True, exist_ok=True)
    input_file_path(job.id).write_bytes(_docx_bytes())
    fake_classifier = FakeClassifier({0: ParagraphRole.HEADING_1, 1: ParagraphRole.BODY})

    await process_job(job.id, classifier=fake_classifier, session_maker=make_test_session)

    async with make_test_session() as session:
        refreshed = await session.get(Job, job.id)
        assert refreshed.status == JobStatus.DONE
        assert refreshed.error_message is None


async def test_reprocessing_an_already_done_job_replaces_its_report():
    job = await _create_job()
    input_file_path(job.id).parent.mkdir(parents=True, exist_ok=True)
    input_file_path(job.id).write_bytes(_docx_bytes())
    fake_classifier = FakeClassifier({0: ParagraphRole.HEADING_1, 1: ParagraphRole.BODY})

    await process_job(job.id, classifier=fake_classifier, session_maker=make_test_session)
    # reprocess the same already-done job - must not crash on the unique
    # constraint on validation_reports.job_id
    await process_job(job.id, classifier=fake_classifier, session_maker=make_test_session)

    async with make_test_session() as session:
        refreshed = await session.get(Job, job.id)
        assert refreshed.status == JobStatus.DONE

        reports = (
            await session.scalars(
                select(ValidationReport).where(ValidationReport.job_id == job.id)
            )
        ).all()
        assert len(reports) == 1
