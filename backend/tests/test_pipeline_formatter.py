from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.pipeline.formatter import apply_formatting
from app.pipeline.rules import FormattingRules
from app.pipeline.schemas import ClassifiedParagraph, ParagraphRole

RULES = FormattingRules(generate_toc=False)


def _build_input(tmp_path: Path) -> Path:
    document = Document()
    document.add_heading("Introduction", level=1)
    document.add_paragraph("Body paragraph in Calibri.")
    path = tmp_path / "input.docx"
    document.save(str(path))
    return path


def test_apply_formatting_sets_margins_font_and_line_spacing(tmp_path: Path):
    input_path = _build_input(tmp_path)
    output_path = tmp_path / "output.docx"
    classified = [
        ClassifiedParagraph(index=0, text="Introduction", role=ParagraphRole.HEADING_1),
        ClassifiedParagraph(index=1, text="Body paragraph in Calibri.", role=ParagraphRole.BODY),
    ]

    changes = apply_formatting(input_path, output_path, classified, RULES)

    assert output_path.exists()
    assert len(changes) == 4  # margins, font/alignment, indent, heading style

    result = Document(str(output_path))
    section = result.sections[0]
    assert round(section.top_margin.mm) == 20
    assert round(section.left_margin.mm) == 30
    assert round(section.right_margin.mm) == 15

    heading_run = result.paragraphs[0].runs[0]
    assert heading_run.font.name == "Times New Roman"
    assert heading_run.font.size.pt == 16  # 14 + full heading_size_bump_pt of 2
    assert heading_run.font.bold is True

    body_paragraph = result.paragraphs[1]
    assert body_paragraph.paragraph_format.line_spacing == 1.5
    body_run = body_paragraph.runs[0]
    assert body_run.font.name == "Times New Roman"
    assert body_run.font.size.pt == 14
    assert body_run.font.bold is not True


def test_apply_formatting_defaults_unclassified_paragraphs_to_body(tmp_path: Path):
    input_path = _build_input(tmp_path)
    output_path = tmp_path / "output.docx"

    # no classification supplied for either paragraph
    apply_formatting(input_path, output_path, classified=[], rules=RULES)

    result = Document(str(output_path))
    for paragraph in result.paragraphs:
        for run in paragraph.runs:
            assert run.font.size.pt == 14
            assert run.font.bold is not True


def test_apply_formatting_sets_cyrillic_capable_font_attributes(tmp_path: Path):
    document = Document()
    document.add_paragraph("Привет, Formatly!")
    input_path = tmp_path / "input.docx"
    document.save(str(input_path))
    output_path = tmp_path / "output.docx"

    apply_formatting(input_path, output_path, classified=[], rules=RULES)

    result = Document(str(output_path))
    run = result.paragraphs[0].runs[0]
    rpr = run._element.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr"
    )
    rfonts = rpr.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts"
    )
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    assert rfonts.get(f"{ns}eastAsia") == "Times New Roman"
    assert rfonts.get(f"{ns}hAnsi") == "Times New Roman"


def test_center_headings_toggle(tmp_path: Path):
    input_path = _build_input(tmp_path)
    classified = [ClassifiedParagraph(index=0, text="Introduction", role=ParagraphRole.HEADING_1)]

    centered_path = tmp_path / "centered.docx"
    apply_formatting(
        input_path,
        centered_path,
        classified,
        FormattingRules(center_headings=True, generate_toc=False),
    )
    centered = Document(str(centered_path)).paragraphs[0]
    assert centered.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER

    left_path = tmp_path / "left.docx"
    apply_formatting(
        input_path,
        left_path,
        classified,
        FormattingRules(center_headings=False, generate_toc=False),
    )
    left = Document(str(left_path)).paragraphs[0]
    assert left.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.LEFT


def test_bold_and_italic_heading_toggles(tmp_path: Path):
    input_path = _build_input(tmp_path)
    output_path = tmp_path / "output.docx"
    classified = [ClassifiedParagraph(index=0, text="Introduction", role=ParagraphRole.HEADING_1)]

    apply_formatting(
        input_path,
        output_path,
        classified,
        FormattingRules(bold_headings=False, italic_headings=True, generate_toc=False),
    )

    run = Document(str(output_path)).paragraphs[0].runs[0]
    assert run.font.bold is not True
    assert run.font.italic is True


def test_heading_size_bump_scales_down_by_level(tmp_path: Path):
    document = Document()
    document.add_heading("H1", level=1)
    document.add_heading("H2", level=2)
    document.add_heading("H3", level=3)
    input_path = tmp_path / "input.docx"
    document.save(str(input_path))
    output_path = tmp_path / "output.docx"

    classified = [
        ClassifiedParagraph(index=0, text="H1", role=ParagraphRole.HEADING_1),
        ClassifiedParagraph(index=1, text="H2", role=ParagraphRole.HEADING_2),
        ClassifiedParagraph(index=2, text="H3", role=ParagraphRole.HEADING_3),
    ]

    apply_formatting(
        input_path,
        output_path,
        classified,
        FormattingRules(font_size_pt=14, heading_size_bump_pt=4, generate_toc=False),
    )

    result = Document(str(output_path))
    assert result.paragraphs[0].runs[0].font.size.pt == 18  # +4, the full bump
    assert result.paragraphs[1].runs[0].font.size.pt == 16  # +2, half the bump
    assert result.paragraphs[2].runs[0].font.size.pt == 14  # +0, no bump


def test_page_break_before_heading_1_toggle(tmp_path: Path):
    input_path = _build_input(tmp_path)
    output_path = tmp_path / "output.docx"
    classified = [ClassifiedParagraph(index=0, text="Introduction", role=ParagraphRole.HEADING_1)]

    apply_formatting(
        input_path,
        output_path,
        classified,
        FormattingRules(page_break_before_heading_1=True, generate_toc=False),
    )

    heading = Document(str(output_path)).paragraphs[0]
    assert heading.paragraph_format.page_break_before is True


def test_paragraph_alignment_applies_to_body_not_headings(tmp_path: Path):
    input_path = _build_input(tmp_path)
    output_path = tmp_path / "output.docx"
    classified = [
        ClassifiedParagraph(index=0, text="Introduction", role=ParagraphRole.HEADING_1),
        ClassifiedParagraph(index=1, text="Body paragraph in Calibri.", role=ParagraphRole.BODY),
    ]

    apply_formatting(
        input_path,
        output_path,
        classified,
        FormattingRules(paragraph_alignment="left", center_headings=True, generate_toc=False),
    )

    result = Document(str(output_path))
    assert result.paragraphs[1].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert result.paragraphs[0].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_paragraph_indent_toggle(tmp_path: Path):
    input_path = _build_input(tmp_path)
    classified = [
        ClassifiedParagraph(index=0, text="Introduction", role=ParagraphRole.HEADING_1),
        ClassifiedParagraph(index=1, text="Body paragraph in Calibri.", role=ParagraphRole.BODY),
    ]

    indented_path = tmp_path / "indented.docx"
    apply_formatting(
        input_path,
        indented_path,
        classified,
        FormattingRules(paragraph_indent_enabled=True, paragraph_indent_mm=12.5, generate_toc=False),
    )
    indented_body = Document(str(indented_path)).paragraphs[1]
    # Word stores indents in twips, not EMU - round-tripping through a saved
    # .docx always introduces a fraction-of-a-mm of noise, so 12.5 comes back
    # as ~12.506 rather than exactly 12.5.
    assert indented_body.paragraph_format.first_line_indent.mm == pytest.approx(12.5, abs=0.1)
    # headings never get a first-line indent regardless of the setting
    indented_heading = Document(str(indented_path)).paragraphs[0]
    assert indented_heading.paragraph_format.first_line_indent.mm == 0

    flush_path = tmp_path / "flush.docx"
    apply_formatting(
        input_path,
        flush_path,
        classified,
        FormattingRules(paragraph_indent_enabled=False, generate_toc=False),
    )
    flush_body = Document(str(flush_path)).paragraphs[1]
    assert flush_body.paragraph_format.first_line_indent.mm == 0


def test_generate_toc_inserts_heading_and_field_before_first_paragraph(tmp_path: Path):
    input_path = _build_input(tmp_path)
    output_path = tmp_path / "output.docx"
    classified = [
        ClassifiedParagraph(index=0, text="Introduction", role=ParagraphRole.HEADING_1),
        ClassifiedParagraph(index=1, text="Body paragraph in Calibri.", role=ParagraphRole.BODY),
    ]

    changes = apply_formatting(
        input_path, output_path, classified, FormattingRules(generate_toc=True)
    )

    assert any("table of contents" in change for change in changes)

    result = Document(str(output_path))
    paragraph_texts = [p.text for p in result.paragraphs]
    assert paragraph_texts[0] == "СОДЕРЖАНИЕ"
    assert paragraph_texts[2:] == ["Introduction", "Body paragraph in Calibri."]

    field_paragraph = result.paragraphs[1]
    instr_texts = [instr.text for instr in field_paragraph._p.iter(qn("w:instrText"))]
    assert any("TOC" in text for text in instr_texts)

    # the paragraph right after the TOC starts a new page
    assert result.paragraphs[2].paragraph_format.page_break_before is True

    update_fields = result.settings.element.find(qn("w:updateFields"))
    assert update_fields is not None
    assert update_fields.get(qn("w:val")) == "true"


def test_generate_toc_inserts_after_title_when_present(tmp_path: Path):
    document = Document()
    document.add_paragraph("Курсовая работа")
    document.add_heading("Introduction", level=1)
    document.add_paragraph("Body paragraph.")
    input_path = tmp_path / "input.docx"
    document.save(str(input_path))
    output_path = tmp_path / "output.docx"

    classified = [
        ClassifiedParagraph(index=0, text="Курсовая работа", role=ParagraphRole.TITLE),
        ClassifiedParagraph(index=1, text="Introduction", role=ParagraphRole.HEADING_1),
        ClassifiedParagraph(index=2, text="Body paragraph.", role=ParagraphRole.BODY),
    ]

    apply_formatting(input_path, output_path, classified, FormattingRules(generate_toc=True))

    result = Document(str(output_path))
    assert [p.text for p in result.paragraphs][:2] == ["Курсовая работа", "СОДЕРЖАНИЕ"]
    assert result.paragraphs[3].text == "Introduction"
    assert result.paragraphs[3].paragraph_format.page_break_before is True


def test_generate_toc_skipped_when_disabled(tmp_path: Path):
    input_path = _build_input(tmp_path)
    output_path = tmp_path / "output.docx"
    classified = [ClassifiedParagraph(index=0, text="Introduction", role=ParagraphRole.HEADING_1)]

    changes = apply_formatting(
        input_path, output_path, classified, FormattingRules(generate_toc=False)
    )

    assert not any("table of contents" in change for change in changes)
    result = Document(str(output_path))
    assert [p.text for p in result.paragraphs] == ["Introduction", "Body paragraph in Calibri."]


def test_generate_toc_skipped_when_no_headings(tmp_path: Path):
    document = Document()
    document.add_paragraph("Just a body paragraph.")
    input_path = tmp_path / "input.docx"
    document.save(str(input_path))
    output_path = tmp_path / "output.docx"

    classified = [
        ClassifiedParagraph(index=0, text="Just a body paragraph.", role=ParagraphRole.BODY)
    ]
    changes = apply_formatting(
        input_path, output_path, classified, FormattingRules(generate_toc=True)
    )

    assert not any("table of contents" in change for change in changes)
    result = Document(str(output_path))
    assert [p.text for p in result.paragraphs] == ["Just a body paragraph."]


def test_completion_is_appended_to_the_paragraphs_text(tmp_path: Path):
    document = Document()
    document.add_paragraph("The results show a clear trend and")
    input_path = tmp_path / "input.docx"
    document.save(str(input_path))
    output_path = tmp_path / "output.docx"

    classified = [
        ClassifiedParagraph(
            index=0,
            text="The results show a clear trend and",
            role=ParagraphRole.BODY,
            completion="continue to improve over time.",
        )
    ]
    changes = apply_formatting(
        input_path, output_path, classified, FormattingRules(generate_toc=False)
    )

    result = Document(str(output_path))
    assert result.paragraphs[0].text == (
        "The results show a clear trend and continue to improve over time."
    )
    # the appended run must pick up the same formatting as the rest of the paragraph
    for run in result.paragraphs[0].runs:
        assert run.font.name == "Times New Roman"
        assert run.font.size.pt == 14
    assert any(
        'completed a paragraph ending "The results show a clear trend and"'
        ' by appending: "continue to improve over time."' in change
        for change in changes
    )


def test_completion_skipped_when_ai_light_editing_disabled(tmp_path: Path):
    document = Document()
    document.add_paragraph("The results show a clear trend and")
    input_path = tmp_path / "input.docx"
    document.save(str(input_path))
    output_path = tmp_path / "output.docx"

    classified = [
        ClassifiedParagraph(
            index=0,
            text="The results show a clear trend and",
            role=ParagraphRole.BODY,
            completion="continue to improve over time.",
        )
    ]
    changes = apply_formatting(
        input_path,
        output_path,
        classified,
        FormattingRules(generate_toc=False, ai_light_editing_enabled=False),
    )

    result = Document(str(output_path))
    assert result.paragraphs[0].text == "The results show a clear trend and"
    assert not any("completed" in change for change in changes)


def test_completion_description_truncates_a_long_original_ending(tmp_path: Path):
    long_text = (
        "This is a much longer paragraph that goes on for a while before it "
        "gets cut off and"
    )
    document = Document()
    document.add_paragraph(long_text)
    input_path = tmp_path / "input.docx"
    document.save(str(input_path))
    output_path = tmp_path / "output.docx"

    classified = [
        ClassifiedParagraph(
            index=0, text=long_text, role=ParagraphRole.BODY, completion="stops there."
        )
    ]
    changes = apply_formatting(
        input_path, output_path, classified, FormattingRules(generate_toc=False)
    )

    completion_change = next(change for change in changes if "appending" in change)
    assert "…" in completion_change
    assert "gets cut off and" in completion_change
    assert "This is a much longer" not in completion_change


def test_completion_description_lists_each_completion_separately(tmp_path: Path):
    document = Document()
    document.add_paragraph("First cut off sentence and")
    document.add_paragraph("Second cut off sentence and")
    input_path = tmp_path / "input.docx"
    document.save(str(input_path))
    output_path = tmp_path / "output.docx"

    classified = [
        ClassifiedParagraph(
            index=0,
            text="First cut off sentence and",
            role=ParagraphRole.BODY,
            completion="so it continues.",
        ),
        ClassifiedParagraph(
            index=1,
            text="Second cut off sentence and",
            role=ParagraphRole.BODY,
            completion="finishes differently.",
        ),
    ]
    changes = apply_formatting(
        input_path, output_path, classified, FormattingRules(generate_toc=False)
    )

    completion_changes = [change for change in changes if "appending" in change]
    assert len(completion_changes) == 2
    assert any("so it continues." in change for change in completion_changes)
    assert any("finishes differently." in change for change in completion_changes)


def test_generated_title_is_inserted_when_none_classified(tmp_path: Path):
    document = Document()
    document.add_heading("Introduction", level=1)
    document.add_paragraph("Body paragraph.")
    input_path = tmp_path / "input.docx"
    document.save(str(input_path))
    output_path = tmp_path / "output.docx"

    classified = [
        ClassifiedParagraph(index=0, text="Introduction", role=ParagraphRole.HEADING_1),
        ClassifiedParagraph(index=1, text="Body paragraph.", role=ParagraphRole.BODY),
    ]
    changes = apply_formatting(
        input_path,
        output_path,
        classified,
        FormattingRules(generate_toc=False),
        generated_title="Курсовая работа",
    )

    result = Document(str(output_path))
    assert [p.text for p in result.paragraphs] == [
        "Курсовая работа",
        "Introduction",
        "Body paragraph.",
    ]
    title_run = result.paragraphs[0].runs[0]
    assert title_run.font.bold is True
    assert result.paragraphs[0].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert any("Курсовая работа" in change for change in changes)


def test_generated_title_skipped_when_a_title_already_classified(tmp_path: Path):
    document = Document()
    document.add_paragraph("Existing Title")
    document.add_paragraph("Body paragraph.")
    input_path = tmp_path / "input.docx"
    document.save(str(input_path))
    output_path = tmp_path / "output.docx"

    classified = [
        ClassifiedParagraph(index=0, text="Existing Title", role=ParagraphRole.TITLE),
        ClassifiedParagraph(index=1, text="Body paragraph.", role=ParagraphRole.BODY),
    ]
    changes = apply_formatting(
        input_path,
        output_path,
        classified,
        FormattingRules(generate_toc=False),
        generated_title="A different generated title",
    )

    result = Document(str(output_path))
    assert [p.text for p in result.paragraphs] == ["Existing Title", "Body paragraph."]
    assert not any("generated a document title" in change for change in changes)


def test_generated_title_skipped_when_ai_light_editing_disabled(tmp_path: Path):
    document = Document()
    document.add_paragraph("Body paragraph.")
    input_path = tmp_path / "input.docx"
    document.save(str(input_path))
    output_path = tmp_path / "output.docx"

    classified = [ClassifiedParagraph(index=0, text="Body paragraph.", role=ParagraphRole.BODY)]
    changes = apply_formatting(
        input_path,
        output_path,
        classified,
        FormattingRules(generate_toc=False, ai_light_editing_enabled=False),
        generated_title="Курсовая работа",
    )

    result = Document(str(output_path))
    assert [p.text for p in result.paragraphs] == ["Body paragraph."]
    assert not any("generated a document title" in change for change in changes)


def test_generated_title_anchors_the_toc_right_after_it(tmp_path: Path):
    document = Document()
    document.add_heading("Introduction", level=1)
    document.add_paragraph("Body paragraph.")
    input_path = tmp_path / "input.docx"
    document.save(str(input_path))
    output_path = tmp_path / "output.docx"

    classified = [
        ClassifiedParagraph(index=0, text="Introduction", role=ParagraphRole.HEADING_1),
        ClassifiedParagraph(index=1, text="Body paragraph.", role=ParagraphRole.BODY),
    ]
    apply_formatting(
        input_path,
        output_path,
        classified,
        FormattingRules(generate_toc=True),
        generated_title="Курсовая работа",
    )

    result = Document(str(output_path))
    paragraph_texts = [p.text for p in result.paragraphs]
    assert paragraph_texts[0] == "Курсовая работа"
    assert paragraph_texts[1] == "СОДЕРЖАНИЕ"
    assert paragraph_texts[3:] == ["Introduction", "Body paragraph."]
    assert result.paragraphs[3].paragraph_format.page_break_before is True
