from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt

from app.pipeline.formatter import apply_formatting
from app.pipeline.rules import FormattingRules
from app.pipeline.schemas import ClassifiedParagraph, ParagraphRole
from app.pipeline.validator import validate_document

RULES = FormattingRules()


def _apply_margins(document: Document, rules: FormattingRules) -> None:
    for section in document.sections:
        section.top_margin = Mm(rules.margins_mm.top)
        section.bottom_margin = Mm(rules.margins_mm.bottom)
        section.left_margin = Mm(rules.margins_mm.left)
        section.right_margin = Mm(rules.margins_mm.right)


def test_validate_document_reports_no_issues_after_formatter_ran(tmp_path: Path):
    document = Document()
    document.add_heading("Introduction", level=1)
    document.add_paragraph("Body paragraph.")
    input_path = tmp_path / "input.docx"
    document.save(str(input_path))

    output_path = tmp_path / "output.docx"
    classified = [
        ClassifiedParagraph(index=0, text="Introduction", role=ParagraphRole.HEADING_1),
        ClassifiedParagraph(index=1, text="Body paragraph.", role=ParagraphRole.BODY),
    ]
    apply_formatting(input_path, output_path, classified, RULES)

    issues = validate_document(output_path, classified, RULES)

    assert issues == []


def test_validate_document_flags_wrong_margins(tmp_path: Path):
    document = Document()
    document.add_paragraph("Body paragraph.")
    for section in document.sections:
        section.top_margin = Mm(10)
    path = tmp_path / "wrong_margins.docx"
    document.save(str(path))

    issues = validate_document(path, classified=[], rules=RULES)

    assert any("top margin" in issue for issue in issues)


def test_validate_document_flags_wrong_body_font_and_size(tmp_path: Path):
    document = Document()
    paragraph = document.add_paragraph("Body paragraph.")
    _apply_margins(document, RULES)
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Mm(RULES.paragraph_indent_mm)
    run = paragraph.runs[0]
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    path = tmp_path / "wrong_font.docx"
    document.save(str(path))

    classified = [ClassifiedParagraph(index=0, text="Body paragraph.", role=ParagraphRole.BODY)]
    issues = validate_document(path, classified, RULES)

    assert any("font is 'Calibri'" in issue for issue in issues)
    assert any("size is 11" in issue for issue in issues)


def test_validate_document_flags_wrong_alignment(tmp_path: Path):
    document = Document()
    paragraph = document.add_paragraph("Body paragraph.")
    _apply_margins(document, RULES)
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # RULES expects justify
    paragraph.paragraph_format.first_line_indent = Mm(RULES.paragraph_indent_mm)
    path = tmp_path / "wrong_alignment.docx"
    document.save(str(path))

    classified = [ClassifiedParagraph(index=0, text="Body paragraph.", role=ParagraphRole.BODY)]
    issues = validate_document(path, classified, RULES)

    assert any("alignment" in issue for issue in issues)


def test_validate_document_flags_missing_paragraph_indent(tmp_path: Path):
    document = Document()
    paragraph = document.add_paragraph("Body paragraph.")
    _apply_margins(document, RULES)
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # no first_line_indent set, but RULES.paragraph_indent_enabled defaults to True
    path = tmp_path / "no_indent.docx"
    document.save(str(path))

    classified = [ClassifiedParagraph(index=0, text="Body paragraph.", role=ParagraphRole.BODY)]
    issues = validate_document(path, classified, RULES)

    assert any("first-line indent" in issue for issue in issues)


def test_validate_document_skips_indent_check_when_disabled(tmp_path: Path):
    document = Document()
    paragraph = document.add_paragraph("Body paragraph.")
    rules = FormattingRules(paragraph_indent_enabled=False)
    _apply_margins(document, rules)
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.runs[0].font.name = rules.font_family
    paragraph.runs[0].font.size = Pt(rules.font_size_pt)
    # deliberately no indent set either way - shouldn't matter since disabled
    path = tmp_path / "indent_disabled.docx"
    document.save(str(path))

    classified = [ClassifiedParagraph(index=0, text="Body paragraph.", role=ParagraphRole.BODY)]
    issues = validate_document(path, classified, rules)

    assert issues == []


def test_validate_document_ignores_heading_paragraphs_for_body_checks(tmp_path: Path):
    document = Document()
    paragraph = document.add_heading("Introduction", level=1)
    _apply_margins(document, RULES)
    paragraph.runs[0].font.name = "Calibri"
    path = tmp_path / "heading_only.docx"
    document.save(str(path))

    classified = [ClassifiedParagraph(index=0, text="Introduction", role=ParagraphRole.HEADING_1)]
    issues = validate_document(path, classified, RULES)

    assert issues == []
